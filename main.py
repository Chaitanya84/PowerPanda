"""
PowerPanda — FastAPI Production Backend
"""

import os
import hashlib
import json
import shutil
import tempfile
import secrets
from pathlib import Path

import numpy as np
import pandas as pd

from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.requests import Request
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument

from extract_graph import extract_from_documents
from build_graph import build_knowledge_graph, get_graph
from query_graphrag import query_powerpanda

import logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(name)s] %(message)s",
    datefmt="%H:%M:%S"
)

load_dotenv()

# ── Config ────────────────────────────────────────────────────────────────────
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")

OPENAI_EMBED_MODEL = os.getenv("OPENAI_EMBED_MODEL", "text-embedding-ada-002")
STORE_DIR      = Path("powerpanda_store")
STORE_DIR.mkdir(exist_ok=True)
EMBEDDED_FILES_RECORD = STORE_DIR / "embedded_files.json"

# ── App ───────────────────────────────────────────────────────────────────────
app = FastAPI(title="OraKa", version="1.0.0")
app.mount("/static", StaticFiles(directory="static"), name="static")
app.mount("/assets", StaticFiles(directory="."), name="assets")   # serves Logo.png
TEMPLATE_DIR = "templates" if Path("templates").exists() else "template"
templates = Jinja2Templates(directory=TEMPLATE_DIR)

# ── In-process caches (single-worker dev mode) ────────────────────────────────
_node_emb_cache: dict = {}
_vectorstore: FAISS | None = None


# ── Helpers ───────────────────────────────────────────────────────────────────

def _md5(data: bytes) -> str:
    return hashlib.md5(data).hexdigest()


def _load_embedded_files() -> dict:
    if EMBEDDED_FILES_RECORD.exists():
        return json.loads(EMBEDDED_FILES_RECORD.read_text())
    return {}


def _save_embedded_files(record: dict):
    EMBEDDED_FILES_RECORD.write_text(json.dumps(record, indent=2))


def _get_embeddings(api_key: str) -> OpenAIEmbeddings:
    return OpenAIEmbeddings(
        model=OPENAI_EMBED_MODEL,
        api_key=api_key,
    )


def _load_vectorstore(api_key: str) -> FAISS | None:
    faiss_path = STORE_DIR / "powerpanda_faiss"
    if faiss_path.exists():
        return FAISS.load_local(
            str(faiss_path),
            _get_embeddings(api_key),
            allow_dangerous_deserialization=True,
        )
    return None


def _save_vectorstore(vs: FAISS, api_key: str):
    faiss_path = STORE_DIR / "powerpanda_faiss"
    vs.save_local(str(faiss_path))


def _parse_file(tmp_path: str, filename: str) -> list[Document]:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return PyPDFLoader(tmp_path).load()
    if ext == "docx":
        doc = DocxDocument(tmp_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(page_content=text, metadata={"source_file": filename})]
    if ext in ("xlsx", "xls"):
        df = pd.read_excel(tmp_path)
        return [Document(page_content=df.to_string(index=False), metadata={"source_file": filename})]
    if ext == "csv":
        df = pd.read_csv(tmp_path)
        return [Document(page_content=df.to_string(index=False), metadata={"source_file": filename})]
    # txt / fallback
    text = Path(tmp_path).read_text(encoding="utf-8", errors="ignore")
    return [Document(page_content=text, metadata={"source_file": filename})]


# ── Auth configuration (env-driven) ───────────────────────────────────────────
AUTH_USERNAME = (
    os.getenv("POWERPANDA_USERNAME")
    or os.getenv("APP_USERNAME")
    or os.getenv("LOGIN_USERNAME")
    or "admin"
)
AUTH_PASSWORD = (
    os.getenv("POWERPANDA_PASSWORD")
    or os.getenv("APP_PASSWORD")
    or os.getenv("LOGIN_PASSWORD")
    or "powerpanda2026"
)

if not (os.getenv("POWERPANDA_USERNAME") or os.getenv("APP_USERNAME") or os.getenv("LOGIN_USERNAME")):
    logging.warning("No username env var found. Falling back to default username.")
if not (os.getenv("POWERPANDA_PASSWORD") or os.getenv("APP_PASSWORD") or os.getenv("LOGIN_PASSWORD")):
    logging.warning("No password env var found. Falling back to default password.")

# ── Routes ────────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def login_page(request: Request):
    """Landing page — login screen."""
    return templates.TemplateResponse(request=request, name="login.html")


@app.post("/api/login")
async def login(request: Request):
    """Authenticate user."""
    body = await request.json()
    username = body.get("username", "").strip()
    password = body.get("password", "")
    if secrets.compare_digest(username, AUTH_USERNAME) and secrets.compare_digest(password, AUTH_PASSWORD):
        return {"token": "authenticated", "user": username}
    raise HTTPException(status_code=401, detail="Invalid username or password.")


@app.get("/app", response_class=HTMLResponse)
async def main_app(request: Request):
    """Main application page (after login)."""
    return templates.TemplateResponse(request=request, name="index.html")


@app.get("/graph-viewer", response_class=HTMLResponse)
async def graph_viewer(request: Request):
    """Serve the interactive graph viewer page."""
    return templates.TemplateResponse(request=request, name="graph_viewer.html")

@app.get("/api/files")
async def list_files():
    """Return already-embedded files."""
    return _load_embedded_files()


@app.post("/api/upload")
async def upload_files(
    files: list[UploadFile] = File(...),
    api_key: str = Form(""),
    graph_name: str = Form("powerpanda"),
):
    """Process uploaded files: embed into FAISS + extract graph."""
    global _vectorstore, _node_emb_cache

    api_key = api_key or OPENAI_API_KEY
    if not api_key:
        raise HTTPException(status_code=400, detail="API key not set in .env or request")
    embedded = _load_embedded_files()
    splitter  = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    embeddings_model = _get_embeddings(api_key)

    results = []
    new_docs: list[Document] = []
    all_texts: list[str]    = []

    for upload in files:
        raw = await upload.read()
        file_hash = _md5(raw)

        if file_hash in embedded.values():
            results.append({"file": upload.filename, "status": "skipped (already embedded)"})
            print(f"[UPLOAD] Skipping '{upload.filename}' — already embedded.")
            continue

        # Write to temp file
        suffix = "." + upload.filename.rsplit(".", 1)[-1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(raw)
            tmp_path = tmp.name

        try:
            docs = _parse_file(tmp_path, upload.filename)
            for d in docs:
                d.metadata["source_file"] = upload.filename
            chunks = splitter.split_documents(docs)
            new_docs.extend(chunks)
            all_texts.extend(d.page_content for d in docs)
            embedded[upload.filename] = file_hash
            results.append({"file": upload.filename, "status": "processed"})
            print(f"[UPLOAD] Processed '{upload.filename}': {len(chunks)} chunks")
        finally:
            os.unlink(tmp_path)

    # Update FAISS
    if new_docs:
        if _vectorstore is None:
            _vectorstore = _load_vectorstore(api_key)

        if _vectorstore:
            _vectorstore.add_documents(new_docs)
        else:
            _vectorstore = FAISS.from_documents(new_docs, embeddings_model)
        _save_vectorstore(_vectorstore, api_key)

        # Extract entities/relations and rebuild graph
        entities, relations = extract_from_documents(all_texts, api_key)
        build_knowledge_graph(entities, relations, graph_name)
        _node_emb_cache = {}  # invalidate node embedding cache

        _save_embedded_files(embedded)

    return {"results": results}


@app.post("/api/query")
async def query(
    request: Request,
):
    """Run the PowerPanda query pipeline."""
    global _vectorstore, _node_emb_cache

    body = await request.json()
    query_text: str = body.get("query", "").strip()
    api_key:    str = body.get("api_key", "") or OPENAI_API_KEY
    graph_name: str = body.get("graph_name", "powerpanda")
    top_k_docs: int = body.get("top_k_docs", 4)
    top_k_nodes:int = body.get("top_k_nodes", 5)
    hops:       int = body.get("hops", 2)

    if not query_text:
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    if not api_key:
        raise HTTPException(status_code=400, detail="API key required")

    if _vectorstore is None:
        _vectorstore = _load_vectorstore(api_key)

    answer, source_docs, graph_context, relevant_nodes = query_powerpanda(
        query=query_text,
        api_key=api_key,
        vectorstore=_vectorstore,
        graph_name=graph_name,
        node_embeddings_cache=_node_emb_cache,
        top_k_docs=top_k_docs,
        top_k_nodes=top_k_nodes,
        hops=hops,
    )

    sources = list({d.metadata.get("source_file", "unknown") for d in source_docs})

    return {
        "answer": answer,
        "sources": sources,
        "graph_context": graph_context,
        "relevant_nodes": relevant_nodes,
    }


@app.get("/api/graph")
async def get_graph_diagram(graph_name: str = "powerpanda", highlight: str = ""):
    """Return Mermaid diagram string for the current graph."""
    graph = get_graph(graph_name)
    highlighted = [n.strip() for n in highlight.split(",") if n.strip()]
    mermaid_code = graph.to_mermaid(highlight_nodes=highlighted)
    node_count = len(graph.nodes)
    edge_count = len(graph.edges)
    return {
        "mermaid": mermaid_code,
        "node_count": node_count,
        "edge_count": edge_count,
    }


@app.delete("/api/clear")
async def clear_data(graph_name: str = "powerpanda"):
    """Delete all embeddings, graph, and records from disk."""
    global _vectorstore, _node_emb_cache

    faiss_path = STORE_DIR / "powerpanda_faiss"
    graph_path = STORE_DIR / f"{graph_name}_graph.json"

    if faiss_path.exists():
        shutil.rmtree(faiss_path)
    if graph_path.exists():
        graph_path.unlink()
    if EMBEDDED_FILES_RECORD.exists():
        EMBEDDED_FILES_RECORD.unlink()

    _vectorstore    = None
    _node_emb_cache = {}

    return {"status": "cleared"}
