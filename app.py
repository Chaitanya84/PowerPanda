"""
Production PowerPanda Streamlit Application

- LLM-based entity/relation extraction (Claude)
- Embedding similarity for node matching (text-embedding-ada-002)
- FAISS vector store for document retrieval
- In-memory knowledge graph with Mermaid visualization
- Claude Opus for answer generation
- Multi-format file upload (PDF, DOCX, XLSX, CSV)
"""

import streamlit as st
import os
import hashlib
import tempfile
import json
from pathlib import Path
import shutil

import pandas as pd
from dotenv import load_dotenv
from langchain_openai import AzureOpenAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from docx import Document as DocxDocument

from extract_graph import extract_from_documents
from build_graph import build_knowledge_graph, get_graph
from query_powerpanda import query_powerpanda

load_dotenv()

AZURE_ENDPOINT = os.getenv("NEXUS_BASE_URL", "https://genai-nexus.api.corpinter.net/")
API_VERSION = "2024-10-21"
EMBEDDINGS_DIR = Path("powerpanda_store")
EMBEDDINGS_DIR.mkdir(exist_ok=True)
EMBEDDED_FILES_RECORD = EMBEDDINGS_DIR / "embedded_files.json"


def get_file_hash(file_content: bytes) -> str:
    """Get MD5 hash of file content for dedup."""
    return hashlib.md5(file_content).hexdigest()


def load_embedded_files_record() -> dict:
    """Load record of already-embedded files."""
    if EMBEDDED_FILES_RECORD.exists():
        return json.loads(EMBEDDED_FILES_RECORD.read_text())
    return {}


def save_embedded_files_record(record: dict):
    """Save record of embedded files."""
    EMBEDDED_FILES_RECORD.write_text(json.dumps(record, indent=2))


def load_docx(file_path: str) -> list[Document]:
    """Load a DOCX file into Documents."""
    doc = DocxDocument(file_path)
    text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return [Document(page_content=text, metadata={"source": file_path})]


def load_xlsx(file_path: str) -> list[Document]:
    """Load an XLSX file into Documents."""
    dfs = pd.read_excel(file_path, sheet_name=None)
    docs = []
    for sheet_name, df in dfs.items():
        text = f"Sheet: {sheet_name}\n{df.to_string(index=False)}"
        docs.append(Document(page_content=text, metadata={"source": file_path, "sheet": sheet_name}))
    return docs


def load_csv(file_path: str) -> list[Document]:
    """Load a CSV file into Documents."""
    df = pd.read_csv(file_path)
    text = df.to_string(index=False)
    return [Document(page_content=text, metadata={"source": file_path})]


def load_txt(file_path: str) -> list[Document]:
    """Load a TXT file into Documents."""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": file_path})]


def load_file(file_path: str, file_type: str) -> list[Document]:
    """Load a file based on its type."""
    if file_type == "pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif file_type == "docx":
        return load_docx(file_path)
    elif file_type == "xlsx":
        return load_xlsx(file_path)
    elif file_type == "csv":
        return load_csv(file_path)
    elif file_type == "txt":
        return load_txt(file_path)
    return []


# ─── Streamlit UI ─────────────────────────────────────────────────────────────

st.set_page_config(page_title="PowerPanda Production", page_icon="🔗", layout="wide")
st.title("🔗 PowerPanda — Production Mode")
st.caption("LLM Extraction • Embedding Similarity • FAISS • Mermaid Graph • Claude Opus")

# Sidebar
api_key = st.sidebar.text_input("Nexus API Key", type="password", value=os.getenv("NEXUS_API_KEY", ""))
graph_name = st.sidebar.text_input("Graph Name", value="powerpanda")
n_hops = st.sidebar.slider("Graph Traversal Hops", 1, 4, 2)
top_k_docs = st.sidebar.slider("Top-K Documents", 1, 10, 4)
top_k_nodes = st.sidebar.slider("Top-K Graph Nodes", 1, 15, 5)

# ─── Clear Data ───────────────────────────────────────────────────────────────
st.sidebar.divider()
if st.sidebar.button("🗑️ Clear All Embeddings & Graph", type="secondary"):
    faiss_path = EMBEDDINGS_DIR / f"{graph_name}_faiss"
    graph_path = EMBEDDINGS_DIR / f"{graph_name}_graph.json"
    if faiss_path.exists():
        shutil.rmtree(faiss_path)
    if graph_path.exists():
        graph_path.unlink()
    if EMBEDDED_FILES_RECORD.exists():
        EMBEDDED_FILES_RECORD.unlink()
    if "node_emb_cache" in st.session_state:
        del st.session_state["node_emb_cache"]
    if "vectorstore_ready" in st.session_state:
        del st.session_state["vectorstore_ready"]
    st.sidebar.success("✅ All data cleared!")
    st.rerun()

if not api_key:
    st.warning("Please enter your Nexus API key in the sidebar.")
    st.stop()

# ─── File Upload ──────────────────────────────────────────────────────────────

st.header("📁 Upload Documents")
uploaded_files = st.file_uploader(
    "Upload PDF, DOCX, XLSX, CSV, or TXT files",
    type=["pdf", "docx", "xlsx", "csv", "txt"],
    accept_multiple_files=True,
)

if uploaded_files:
    # Check which files are already embedded
    record = load_embedded_files_record()
    files_to_process = []
    skipped_files = []

    for uf in uploaded_files:
        content = uf.read()
        uf.seek(0)
        file_hash = get_file_hash(content)
        if file_hash in record:
            skipped_files.append(uf.name)
        else:
            files_to_process.append((uf, file_hash, content))

    if skipped_files:
        st.info(f"⏭️ Skipping already-embedded files: {', '.join(skipped_files)}")

    if files_to_process:
        if st.button(f"🚀 Process {len(files_to_process)} new file(s)", type="primary"):
            embeddings = AzureOpenAIEmbeddings(
                azure_deployment="text-embedding-ada-002",
                azure_endpoint=AZURE_ENDPOINT,
                api_key=api_key,
                api_version=API_VERSION,
            )

            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            all_chunks = []
            all_raw_texts = []

            progress = st.progress(0)
            status = st.status("Processing files...", expanded=True)

            for i, (uf, file_hash, content) in enumerate(files_to_process):
                file_type = uf.name.rsplit(".", 1)[-1].lower()
                status.write(f"📄 Loading {uf.name}...")

                # Save to temp file
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type}") as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name

                # Load and chunk
                documents = load_file(tmp_path, file_type)
                chunks = splitter.split_documents(documents)
                for chunk in chunks:
                    chunk.metadata["source_file"] = uf.name
                all_chunks.extend(chunks)

                # Collect raw text for graph extraction
                raw_text = "\n".join(doc.page_content for doc in documents)
                all_raw_texts.append(raw_text)

                os.unlink(tmp_path)
                progress.progress((i + 1) / len(files_to_process))

            # Step 1: Embed documents into FAISS
            status.write("🔢 Embedding documents into FAISS...")
            faiss_path = str(EMBEDDINGS_DIR / f"{graph_name}_faiss")

            if Path(faiss_path).exists():
                vectorstore = FAISS.load_local(faiss_path, embeddings, allow_dangerous_deserialization=True)
                new_vs = FAISS.from_documents(all_chunks, embeddings)
                vectorstore.merge_from(new_vs)
            else:
                vectorstore = FAISS.from_documents(all_chunks, embeddings)

            vectorstore.save_local(faiss_path)

            # Step 2: Extract entities/relations with LLM
            status.write("🧠 Extracting entities & relations with Claude...")
            entities, relations = extract_from_documents(all_raw_texts, api_key=api_key)
            status.write(f"   Found {len(entities)} entities, {len(relations)} relations")

            # Step 3: Build knowledge graph (in-memory + saved to disk)
            status.write("🗂️ Building knowledge graph...")
            graph = build_knowledge_graph(entities, relations, graph_name=graph_name)
            status.write(f"   ✅ Graph built: {len(graph.nodes)} nodes, {len(graph.edges)} edges")

            # Update record
            for uf, file_hash, _ in files_to_process:
                record[file_hash] = uf.name
            save_embedded_files_record(record)

            status.update(label="✅ Processing complete!", state="complete")
            st.success(f"Processed {len(files_to_process)} file(s): {len(all_chunks)} chunks embedded, {len(entities)} entities extracted.")
            st.session_state["vectorstore_ready"] = True

    elif not files_to_process and skipped_files:
        st.session_state["vectorstore_ready"] = True

# ─── Knowledge Graph Visualization ───────────────────────────────────────────

st.header("🕸️ Knowledge Graph")
graph = get_graph(graph_name)
if graph.nodes:
    mermaid_code = graph.to_mermaid()
    st.markdown(f"```mermaid\n{mermaid_code}\n```")
    with st.expander(f"📊 Graph Stats: {len(graph.nodes)} nodes, {len(graph.edges)} edges"):
        st.write("**Entity Types:**")
        type_counts = {}
        for etype in graph.nodes.values():
            type_counts[etype] = type_counts.get(etype, 0) + 1
        for etype, count in sorted(type_counts.items(), key=lambda x: -x[1]):
            st.write(f"  • {etype}: {count}")
else:
    st.caption("No graph built yet. Upload and process documents first.")

# ─── Query Interface ──────────────────────────────────────────────────────────

st.header("💬 Ask Questions")

# Load vectorstore if available
faiss_path = str(EMBEDDINGS_DIR / f"{graph_name}_faiss")
vectorstore = None
if Path(faiss_path).exists():
    embeddings = AzureOpenAIEmbeddings(
        azure_deployment="text-embedding-ada-002",
        azure_endpoint=AZURE_ENDPOINT,
        api_key=api_key,
        api_version=API_VERSION,
    )
    vectorstore = FAISS.load_local(faiss_path, embeddings, allow_dangerous_deserialization=True)
    st.caption("✅ Vector store loaded")

# Show embedded files
record = load_embedded_files_record()
if record:
    with st.expander(f"📂 Embedded Files ({len(record)})"):
        for fname in record.values():
            st.write(f"• {fname}")

question = st.text_input("Enter your question:")

if question:
    if not vectorstore:
        st.warning("No documents embedded yet. Please upload and process files first.")
    else:
        with st.spinner("🔍 Querying PowerPanda pipeline..."):
            # Initialize node embeddings cache in session state
            if "node_emb_cache" not in st.session_state:
                st.session_state["node_emb_cache"] = {}

            answer, source_docs, graph_context, relevant_nodes = query_powerpanda(
                query=question,
                api_key=api_key,
                vectorstore=vectorstore,
                graph_name=graph_name,
                node_embeddings_cache=st.session_state["node_emb_cache"],
                top_k_docs=top_k_docs,
                top_k_nodes=top_k_nodes,
                hops=n_hops,
            )

        st.markdown("### 📝 Answer")
        st.write(answer)

        col1, col2 = st.columns(2)

        with col1:
            with st.expander("📊 Graph Context (Mermaid)"):
                if relevant_nodes and graph.nodes:
                    # Show subgraph with highlighted relevant nodes
                    mermaid_highlighted = graph.to_mermaid(highlight_nodes=relevant_nodes)
                    st.markdown(f"```mermaid\n{mermaid_highlighted}\n```")
                    st.caption(f"Highlighted nodes: {', '.join(relevant_nodes)}")
                if graph_context:
                    st.code(graph_context, language="text")
                elif not relevant_nodes:
                    st.write("No graph context retrieved.")

        with col2:
            with st.expander("📚 Source Documents"):
                for i, doc in enumerate(source_docs):
                    src_file = doc.metadata.get("source_file", "N/A")
                    st.markdown(f"**Chunk {i+1}** — {src_file}")
                    st.caption(doc.page_content[:400])
                    st.divider()
